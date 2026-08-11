"""Smoke test: build two .docx files via python-docx, glue them with
docxcompose, and read the merged result back.

Covers the python-docx + docxcompose leg of the job toolchain end to end —
authoring, merge, and readback — rather than just importing the packages.
"""

import tempfile
from pathlib import Path

from _common import report
from docx import Document

# docxcompose ships no stubs/py.typed marker — same class of issue as root
# pyproject.toml's pdfkit/mdx_math/fuzzysearch overrides; kept as a local
# ignore since a [[tool.mypy.overrides]] entry would touch a file outside
# T3.7's scope (services/executor/** + Makefile).
from docxcompose.composer import Composer  # type: ignore[import-untyped]

NAME = "python_docx"


def _build_doc(path: Path, heading: str, paragraph: str) -> None:
    doc = Document()
    doc.add_heading(heading, level=1)
    doc.add_paragraph(paragraph)
    doc.save(str(path))


def main() -> None:
    with tempfile.TemporaryDirectory(prefix="smoke-docx-") as tmpdir:
        tmp = Path(tmpdir)
        doc_a = tmp / "a.docx"
        doc_b = tmp / "b.docx"
        merged = tmp / "merged.docx"

        _build_doc(doc_a, "Part A", "smoke paragraph A")
        _build_doc(doc_b, "Part B", "smoke paragraph B")

        composer = Composer(Document(str(doc_a)))
        composer.append(Document(str(doc_b)))
        composer.save(str(merged))

        if not merged.is_file() or merged.stat().st_size == 0:
            raise AssertionError("merged docx was not written")

        result = Document(str(merged))
        text = "\n".join(p.text for p in result.paragraphs)
        for expected in ("Part A", "smoke paragraph A", "Part B", "smoke paragraph B"):
            if expected not in text:
                raise AssertionError(
                    f"{expected!r} missing from merged docx text: {text!r}"
                )


if __name__ == "__main__":
    report(NAME, main)
