#!/usr/bin/env python3
"""Build GOST-formatted lab report from Markdown.

Usage:
    python build.py <report.md> [output.docx]

Expects templates/ and filters/ directories next to scripts/.
"""

import shutil
import subprocess
import sys
import tempfile
import re
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docxcompose.composer import Composer

SKILL_DIR = Path(__file__).resolve().parent.parent
TEMPLATES_DIR = SKILL_DIR / "templates"
FILTERS_DIR = SKILL_DIR / "filters"

REFERENCE_DOC = TEMPLATES_DIR / "reference.docx"
TITLE_PAGES_DIR = TEMPLATES_DIR / "title_pages"
DEFAULT_TITLE_PAGE = "guap_lab"
PAGEBREAK_FILTER = FILTERS_DIR / "pagebreak.lua"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# Section headings that should be rendered as ГОСТ structural elements
# (centered, all caps, no numbering). Match by lowercase exact heading text.
GOST_STRUCTURAL_HEADINGS = {
    "содержание",
    "оглавление",
    "введение",
    "заключение",
    "список использованных источников",
    "список литературы",
    "список сокращений",
    "перечень сокращений",
    "реферат",
    "аннотация",
}


def resolve_title_page(metadata: dict) -> Path:
    """Pick the title page template by the `title_page` YAML key.

    Templates live in templates/title_pages/<name>.docx. Falls back to
    DEFAULT_TITLE_PAGE when the key is absent.
    """
    name = metadata.get("title_page", "").strip() or DEFAULT_TITLE_PAGE
    path = TITLE_PAGES_DIR / f"{name}.docx"
    if not path.exists():
        available = sorted(p.stem for p in TITLE_PAGES_DIR.glob("*.docx"))
        print(f"Error: title page template not found: {path}")
        print(f"Available templates: {', '.join(available) or '(none)'}")
        sys.exit(1)
    return path


def parse_yaml_front_matter(md_path: Path) -> dict:
    text = md_path.read_text(encoding="utf-8-sig")
    match = re.match(r"^---\s*\n(.*?)\n---", text, re.DOTALL)
    if not match:
        return {}
    metadata = {}
    for line in match.group(1).splitlines():
        line = line.strip()
        if ":" in line:
            key, _, value = line.partition(":")
            key = key.strip()
            value = value.strip().strip('"').strip("'")
            metadata[key] = value
    return metadata


def fill_title_page(src_path: Path, metadata: dict, output_path: Path):
    shutil.copy2(src_path, output_path)
    doc = Document(str(output_path))

    # Generic mapping: every YAML key becomes {{KEY}} — title page templates
    # are free to use any fields they need (faculty, programme, topic, ...).
    # {{KEY_UPPER}} gives the uppercased value (регистр решает шаблон, не код).
    placeholders = {}
    for key, value in metadata.items():
        placeholders["{{" + key.upper() + "}}"] = str(value)
        placeholders["{{" + key.upper() + "_UPPER}}"] = str(value).upper()
    placeholders.setdefault("{{REPORT_TYPE}}", "ЛАБОРАТОРНОЙ РАБОТЕ")
    placeholders.setdefault("{{YEAR}}", str(datetime.now().year))

    leftover = re.compile(r"\{\{[A-Z0-9_]+\}\}")

    def replace_in_runs(runs):
        for run in runs:
            for placeholder, value in placeholders.items():
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)
            # Wipe placeholders the metadata didn't cover (keeps output clean)
            for missing in leftover.findall(run.text):
                print(f"Warning: no value for {missing} in YAML, leaving blank")
                run.text = run.text.replace(missing, "")

    for paragraph in doc.paragraphs:
        replace_in_runs(paragraph.runs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_runs(paragraph.runs)

    # Optional: title_label fully overrides the "ОТЧЕТ О ... № ..." line
    title_label = metadata.get("title_label", "").strip()
    if title_label:
        _override_title_line(doc, placeholders, title_label)

    doc.save(str(output_path))


def _override_title_line(doc, placeholders, title_label: str):
    """Replace the "ОТЧЕТ О <type> № <number>" line with `title_label`.

    Works after standard placeholder substitution: searches for paragraphs whose
    full text equals the rendered title line, then collapses them to a single
    run containing `title_label`. Preserves the paragraph's run-level formatting
    by reusing the first run.
    """
    rendered_type = placeholders.get("{{REPORT_TYPE}}", "")
    rendered_num = placeholders.get("{{LAB_NUMBER}}", "")
    candidate_full = f"ОТЧЕТ О {rendered_type} № {rendered_num}".strip().rstrip("№").strip()
    candidate_short = f"ОТЧЕТ О {rendered_type}".strip()

    def all_paragraphs(doc):
        yield from doc.paragraphs
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs

    for paragraph in all_paragraphs(doc):
        full = "".join(run.text for run in paragraph.runs).strip().rstrip("№").strip()
        if full and (full == candidate_full or full == candidate_short):
            if not paragraph.runs:
                continue
            first = paragraph.runs[0]
            first.text = title_label
            for run in paragraph.runs[1:]:
                run.text = ""
            return


def transplant_reference_styles(title_path: Path):
    """Ensure the filled title docx carries the GOST styles from reference.docx.

    docxcompose keeps the master document's (= title's) style definitions on
    name clashes, so a title template from a foreign university docx would
    override body styles (headings, body text). Fix: replace/append style
    definitions and docDefaults from reference.docx into the title document.
    No-op when the title already has identical styles (e.g. guap_lab).
    """
    # Normalize reference.docx through the same python-docx serialization the
    # filled title went through, so byte comparison detects a genuine match.
    with tempfile.TemporaryDirectory() as tmp:
        ref_norm = Path(tmp) / "ref_norm.docx"
        shutil.copy2(REFERENCE_DOC, ref_norm)
        ref_doc = Document(str(ref_norm))
        _ = ref_doc.styles  # force styles part to parse
        ref_doc.save(str(ref_norm))
        with zipfile.ZipFile(ref_norm, "r") as zf:
            ref_styles = zf.read("word/styles.xml").decode("utf-8")

    with zipfile.ZipFile(title_path, "r") as zf:
        names = zf.namelist()
        files = {n: zf.read(n) for n in names}
    tpl_styles = files["word/styles.xml"].decode("utf-8")

    if tpl_styles == ref_styles:
        return

    style_re = re.compile(r"<w:style [^>]*>.*?</w:style>", re.DOTALL)
    id_re = re.compile(r'w:styleId="([^"]+)"')

    ref_blocks = {}
    for block in style_re.findall(ref_styles):
        m = id_re.search(block)
        if m:
            # Drop numPr: reference numIds point into reference numbering.xml,
            # which is NOT transplanted — in the merged doc the same numId
            # would resolve to the body's list numbering (visible numbers).
            block = re.sub(r"<w:numPr>.*?</w:numPr>", "", block, flags=re.DOTALL)
            ref_blocks[m.group(1)] = block

    # Replace same-id styles with the reference definitions
    replaced = set()

    def swap(m: re.Match) -> str:
        sid_m = id_re.search(m.group(0))
        if sid_m and sid_m.group(1) in ref_blocks:
            replaced.add(sid_m.group(1))
            return ref_blocks[sid_m.group(1)]
        return m.group(0)

    tpl_styles = style_re.sub(swap, tpl_styles)

    # Append reference styles the template didn't have. docDefaults stay
    # untouched: reference body styles are fully explicit (Normal included),
    # while title paragraphs must be bound to their own TitleText style —
    # NOT Normal — so the GOST Normal transplant can't reflow the title page.
    # Strip w:default — the title doc keeps its own default styles; body
    # styles are referenced by id (pStyle/basedOn) and don't need the flag.
    missing = "".join(
        block.replace(' w:default="1"', "")
        for sid, block in ref_blocks.items()
        if sid not in replaced
    )
    tpl_styles = tpl_styles.replace("</w:styles>", missing + "</w:styles>")

    files["word/styles.xml"] = tpl_styles.encode("utf-8")
    title_path.unlink()
    with zipfile.ZipFile(title_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for n, data in files.items():
            zout.writestr(n, data)


def preprocess_markdown(md_path: Path, output_path: Path):
    text = md_path.read_text(encoding="utf-8-sig")
    text = text.replace("—", "–")  # em dash -> en dash
    text = re.sub(r"(?m)^title:.*\n", "", text)  # strip title from YAML (already on title page)
    output_path.write_text(text, encoding="utf-8")


def build_report_body(md_path: Path, output_path: Path):
    preprocessed = md_path.parent / f".{md_path.stem}_preprocessed.md"
    preprocess_markdown(md_path, preprocessed)
    try:
        cmd = [
            "pandoc",
            str(preprocessed),
            "-o",
            str(output_path),
            f"--reference-doc={REFERENCE_DOC}",
            f"--lua-filter={PAGEBREAK_FILTER}",
            "--highlight-style=pygments",  # выбор владельца 2026-07-16; фиксируем явно от смены дефолта pandoc
        ]
        subprocess.run(cmd, check=True, cwd=md_path.parent)
    finally:
        preprocessed.unlink(missing_ok=True)


def _postprocess_body(body_path: Path, metadata: dict):
    """Apply ГОСТ-style fixes to the pandoc-generated body docx.

    Operates by editing word/document.xml directly (string-level XML edits) to
    avoid python-docx's narrow API for table layout / fields. Idempotent: safe
    to run multiple times on the same file.
    """
    def flag(name):
        return str(metadata.get(name, "")).strip().lower() in ("true", "yes", "1")

    add_toc = flag("add_toc")
    number_headings = flag("number_headings")
    sections_new_page = flag("sections_new_page")

    with zipfile.ZipFile(body_path, "r") as zf:
        names = zf.namelist()
        document_xml = zf.read("word/document.xml").decode("utf-8")

    document_xml = _fix_tables(document_xml)
    document_xml = _fit_table_widths(document_xml)
    document_xml = _fix_structural_headings(document_xml)
    document_xml = _center_appendix_titles(document_xml)
    if number_headings:
        document_xml = _number_headings(document_xml)
    if sections_new_page:
        document_xml = _sections_new_page(document_xml)
    if add_toc:
        document_xml = _insert_toc_field(document_xml)

    # Repack
    with zipfile.ZipFile(body_path, "r") as zin:
        files = {n: zin.read(n) for n in names}
    files["word/document.xml"] = document_xml.encode("utf-8")
    body_path.unlink()
    with zipfile.ZipFile(body_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for n, data in files.items():
            zout.writestr(n, data)


def _fix_tables(document_xml: str) -> str:
    """For every <w:tbl>:
    - tblLayout fixed → autofit
    - tblW: w=0, type=auto (let Word recalc)
    - first row: paragraphs in cells get jc=center
    - insert empty paragraph after table for spacing
    """

    def process_table(match: re.Match) -> str:
        tbl = match.group(0)

        # tblLayout: fixed -> autofit (handle both self-closing and full forms)
        tbl = re.sub(
            r'<w:tblLayout[^>]*w:type="fixed"[^>]*/>',
            '<w:tblLayout w:type="autofit"/>',
            tbl,
        )
        tbl = re.sub(
            r'<w:tblLayout[^>]*w:type="fixed"[^>]*>\s*</w:tblLayout>',
            '<w:tblLayout w:type="autofit"/>',
            tbl,
        )
        # If table has no tblLayout at all, inject autofit
        if '<w:tblLayout' not in tbl:
            tbl = tbl.replace('</w:tblPr>', '<w:tblLayout w:type="autofit"/></w:tblPr>', 1)

        # tblW: set to auto
        tbl = re.sub(
            r'<w:tblW[^/]*/>',
            '<w:tblW w:w="0" w:type="auto"/>',
            tbl,
            count=1,
        )

        # Drop fixed gridCol widths so columns can autofit
        # tbl = re.sub(r'<w:tblGrid>.*?</w:tblGrid>',
        #              lambda m: re.sub(r'<w:gridCol[^/]*/>', '<w:gridCol/>', m.group(0)),
        #              tbl, flags=re.DOTALL)

        # Center the first row's cell paragraphs
        first_row = re.search(r'<w:tr\b[^>]*>.*?</w:tr>', tbl, re.DOTALL)
        if first_row:
            new_first = _center_row_paragraphs(first_row.group(0))
            tbl = tbl.replace(first_row.group(0), new_first, 1)

        # cantSplit: строка таблицы не режется границей страницы, а целиком
        # переезжает на следующую (иначе на разрыве остаётся огрызок строки)
        def add_cantsplit(row_m: re.Match) -> str:
            row = row_m.group(0)
            if "<w:cantSplit" in row:
                return row
            if "<w:trPr>" in row:
                return row.replace("<w:trPr>", "<w:trPr><w:cantSplit/>", 1)
            return re.sub(r"(<w:tr\b[^>]*>)",
                          r"\1<w:trPr><w:cantSplit/></w:trPr>", row, count=1)

        tbl = re.sub(r"<w:tr\b.*?</w:tr>", add_cantsplit, tbl, flags=re.DOTALL)

        return tbl

    document_xml = re.sub(r'<w:tbl>.*?</w:tbl>', process_table, document_xml, flags=re.DOTALL)

    # Always add an empty body-text paragraph after each table to ensure gap
    spacer = '<w:p><w:pPr><w:pStyle w:val="BodyText"/></w:pPr></w:p>'
    document_xml = re.sub(
        r'</w:tbl>',
        lambda m: m.group(0) + spacer,
        document_xml,
    )
    return document_xml


def _fit_table_widths(document_xml: str) -> str:
    """Пересчёт ширин колонок по фактическому содержимому ячеек.

    pandoc назначает gridCol пропорционально ширине колонок в исходном
    Markdown, а не по содержимому: колонка с короткими данными и длинным
    заголовком получает ширину «по данным», и заголовок ломается посреди
    слова. Минимум колонки — её самое длинное неразрывное слово; остаток
    рабочей области раздаётся пропорционально объёму текста, но колонка,
    целиком помещающаяся в одну строку, шире этой строки не делается.
    Ширины — стартовая раскладка для рендереров, честных к сохранённым
    значениям (LibreOffice/печать); tblLayout остаётся autofit, чтобы Word
    мог пересчитать по своим метрикам.
    """
    CHAR_W = 150   # средняя ширина символа TNR 14pt в twips (с запасом под жирные заголовки)
    PAD = 250      # внутренние поля ячейки (2×108) + рамки
    AVAIL = 9638   # рабочая область A4 при полях слева/справа 20 мм
    MIN_COL = 700

    def cell_text(tc_xml: str) -> str:
        return " ".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", tc_xml))

    def process_table(match: re.Match) -> str:
        tbl = match.group(0)
        # объединённые ячейки и картинки в ячейках — раскладку не трогаем
        if "gridSpan" in tbl or "<w:drawing" in tbl:
            return tbl
        grid = re.search(r"<w:tblGrid>.*?</w:tblGrid>", tbl, re.DOTALL)
        if not grid:
            return tbl
        ncols = len(re.findall(r"<w:gridCol", grid.group(0)))
        rows = re.findall(r"<w:tr\b.*?</w:tr>", tbl, re.DOTALL)
        if not ncols or not rows:
            return tbl
        cols = [[] for _ in range(ncols)]
        for row in rows:
            cells = re.findall(r"<w:tc>.*?</w:tc>", row, re.DOTALL)
            if len(cells) != ncols:
                return tbl
            for j, c in enumerate(cells):
                cols[j].append(cell_text(c))

        min_w, max_w, weight = [], [], []
        for texts in cols:
            words = [w for t in texts for w in t.split()] or [""]
            longest_word = max(len(w) for w in words)
            longest_cell = max(len(t) for t in texts)
            min_w.append(max(MIN_COL, longest_word * CHAR_W + PAD))
            max_w.append(max(MIN_COL, longest_cell * CHAR_W + PAD))
            weight.append(sum(len(t) for t in texts) + 1)

        if sum(min_w) >= AVAIL:
            # даже длиннейшие слова не помещаются — ужимаем пропорционально
            scale = AVAIL / sum(min_w)
            widths = [int(w * scale) for w in min_w]
        else:
            widths = list(min_w)
            extra = AVAIL - sum(min_w)
            for _ in range(3):  # излишки сверх «одной строки» перераздаются
                growable = [j for j in range(ncols) if widths[j] < max_w[j]]
                total_weight = sum(weight[j] for j in growable)
                if not growable or not total_weight or extra <= 0:
                    break
                leftover = 0
                for j in growable:
                    add = int(extra * weight[j] / total_weight)
                    if widths[j] + add > max_w[j]:
                        leftover += widths[j] + add - max_w[j]
                        widths[j] = max_w[j]
                    else:
                        widths[j] += add
                extra = leftover

        total = sum(widths)
        new_grid = ("<w:tblGrid>"
                    + "".join(f'<w:gridCol w:w="{w}"/>' for w in widths)
                    + "</w:tblGrid>")
        tbl = tbl.replace(grid.group(0), new_grid, 1)
        tbl = re.sub(r"<w:tblW[^/]*/>",
                     f'<w:tblW w:w="{total}" w:type="dxa"/>', tbl, count=1)

        def fix_row(m: re.Match) -> str:
            idx = [0]

            def fix_cell(cm: re.Match) -> str:
                w = widths[idx[0]]
                idx[0] += 1
                return re.sub(r"<w:tcW[^/]*/>",
                              f'<w:tcW w:w="{w}" w:type="dxa"/>',
                              cm.group(0), count=1)

            return re.sub(r"<w:tc>.*?</w:tc>", fix_cell, m.group(0),
                          flags=re.DOTALL)

        tbl = re.sub(r"<w:tr\b.*?</w:tr>", fix_row, tbl, flags=re.DOTALL)
        return tbl

    return re.sub(r"<w:tbl>.*?</w:tbl>", process_table, document_xml,
                  flags=re.DOTALL)


def _center_row_paragraphs(row_xml: str) -> str:
    """Inside a <w:tr>, set each <w:p>'s pPr to include <w:jc w:val="center"/>."""

    def fix_para(m: re.Match) -> str:
        p = m.group(0)
        if '<w:pPr>' in p:
            # Replace existing jc or insert new one before any <w:rPr>/end of pPr
            if re.search(r'<w:jc[^/]*/>', p):
                p = re.sub(r'<w:jc[^/]*/>', '<w:jc w:val="center"/>', p)
            else:
                p = p.replace('<w:pPr>', '<w:pPr><w:jc w:val="center"/>', 1)
        else:
            p = p.replace('<w:p>', '<w:p><w:pPr><w:jc w:val="center"/></w:pPr>', 1)
        return p

    return re.sub(r'<w:p\b[^/]*?>(?:(?!</w:p>).)*?</w:p>', fix_para, row_xml, flags=re.DOTALL)


def _fix_structural_headings(document_xml: str) -> str:
    """For headings whose text matches a ГОСТ structural section name, replace
    the heading style with a centered all-caps version, drop heading numbering,
    and uppercase the run text."""

    def process_para(m: re.Match) -> str:
        p = m.group(0)
        style_match = re.search(r'<w:pStyle w:val="(Heading[1-9])"', p)
        if not style_match:
            return p
        # Extract plain text
        texts = re.findall(r'<w:t[^>]*>([^<]*)</w:t>', p)
        plain = "".join(texts).strip().lower()
        is_appendix = re.fullmatch(r"приложение\s+[а-яё]", plain) is not None
        if plain not in GOST_STRUCTURAL_HEADINGS and not is_appendix:
            return p
        # Uppercase all run text
        def up(t):
            return f'<w:t xml:space="preserve">{t.group(1).upper()}</w:t>'
        p = re.sub(r'<w:t[^>]*>([^<]*)</w:t>', up, p)
        # Force centered alignment in pPr
        if re.search(r'<w:jc[^/]*/>', p):
            p = re.sub(r'<w:jc[^/]*/>', '<w:jc w:val="center"/>', p)
        elif '<w:pPr>' in p:
            p = p.replace('<w:pPr>', '<w:pPr><w:jc w:val="center"/>', 1)
        return p

    return re.sub(r'<w:p\b[^/]*?>(?:(?!</w:p>).)*?</w:p>', process_para, document_xml, flags=re.DOTALL)


_PARA_RE = re.compile(r"<w:p\b[^/>]*?>(?:(?!</w:p>).)*?</w:p>", re.DOTALL)


def _para_heading_level(p_xml: str):
    """Return heading level (1-3) for a paragraph, or None."""
    m = re.search(r'<w:pStyle w:val="Heading([1-3])"', p_xml)
    return int(m.group(1)) if m else None


def _para_text(p_xml: str) -> str:
    return "".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", p_xml))


def _is_unnumbered_heading(text: str) -> bool:
    plain = text.strip().lower()
    return plain in GOST_STRUCTURAL_HEADINGS or plain.startswith("приложение ")


def _number_headings(document_xml: str) -> str:
    """Prefix Heading1-3 text with ГОСТ 7.32 numbers («1 Название», «1.1 …»).

    Structural headings (введение, заключение, …) and приложения are skipped.
    Numbers are inserted as plain text so the .docx stays fully static.
    """
    counters = [0, 0, 0]

    def process(m: re.Match) -> str:
        p = m.group(0)
        level = _para_heading_level(p)
        if level is None:
            return p
        text = _para_text(p)
        if not text.strip() or _is_unnumbered_heading(text):
            return p
        if re.match(r"^\s*\d+(\.\d+)*\s", text):
            return p  # автор уже пронумеровал — не дублируем
        counters[level - 1] += 1
        for i in range(level, 3):
            counters[i] = 0
        number = ".".join(str(c) for c in counters[:level])
        return re.sub(
            r"(<w:t[^>]*>)",
            lambda t: t.group(1) + number + " ",
            p,
            count=1,
        )

    return _PARA_RE.sub(process, document_xml)


def _sections_new_page(document_xml: str) -> str:
    """Give every Heading1 a pageBreakBefore (each section starts on a new
    page, ГОСТ 7.32). Skipped for the first paragraph of the body and for
    headings already preceded by an explicit page break (agent's \\newpage)."""
    paras = list(_PARA_RE.finditer(document_xml))
    out = []
    last_end = 0
    prev_p = None
    for m in paras:
        p = m.group(0)
        out.append(document_xml[last_end:m.start()])
        if _para_heading_level(p) == 1 and prev_p is not None:
            prev_is_break = (
                'w:type="page"' in prev_p and not _para_text(prev_p).strip()
            )
            has_break_before = "<w:pageBreakBefore" in p
            if not prev_is_break and not has_break_before:
                if "<w:pPr>" in p:
                    p = p.replace("<w:pPr>", "<w:pPr><w:pageBreakBefore/>", 1)
                else:
                    p = re.sub(
                        r"(<w:p\b[^/>]*?>)",
                        r"\1<w:pPr><w:pageBreakBefore/></w:pPr>",
                        p,
                        count=1,
                    )
        out.append(p)
        last_end = m.end()
        prev_p = m.group(0)
    out.append(document_xml[last_end:])
    return "".join(out)


def _center_appendix_titles(document_xml: str) -> str:
    """ГОСТ 7.32: заголовок приложения — отдельной строкой по центру. Центрирует
    первый обычный параграф сразу после заголовка «ПРИЛОЖЕНИЕ X»."""
    paras = list(_PARA_RE.finditer(document_xml))
    out, last_end = [], 0
    prev_was_appendix = False
    for m in paras:
        p = m.group(0)
        out.append(document_xml[last_end:m.start()])
        if prev_was_appendix and _para_heading_level(p) is None and _para_text(p).strip():
            if re.search(r"<w:jc[^/]*/>", p):
                p = re.sub(r"<w:jc[^/]*/>", '<w:jc w:val="center"/>', p, count=1)
            elif "<w:pPr>" in p:
                p = p.replace("<w:pPr>", '<w:pPr><w:jc w:val="center"/>', 1)
            else:
                p = re.sub(r"(<w:p\b[^/>]*?>)", r'\1<w:pPr><w:jc w:val="center"/></w:pPr>', p, count=1)
            prev_was_appendix = False
        else:
            plain = _para_text(p).strip().lower()
            prev_was_appendix = (
                _para_heading_level(p) is not None
                and re.fullmatch(r"приложение\s+[а-яё]", plain) is not None
            )
        out.append(p)
        last_end = m.end()
    out.append(document_xml[last_end:])
    return "".join(out)


def _insert_toc_field(document_xml: str) -> str:
    """Insert an automatic Word TOC field. Placed at the very start of the
    body, or — when the document opens with РЕФЕРАТ/АННОТАЦИЯ (ВКР) — right
    after that section, per ГОСТ 7.32 element order.
    """
    if 'TOC \\o' in document_xml:
        return document_xml  # already inserted

    toc_core = (
        # Heading line
        '<w:p><w:pPr><w:pStyle w:val="UnnumberedHeading1NoTOC1"/>'
        '<w:jc w:val="center"/></w:pPr>'
        '<w:r><w:t xml:space="preserve">СОДЕРЖАНИЕ</w:t></w:r></w:p>'
        # TOC field
        '<w:p><w:pPr><w:pStyle w:val="BodyText"/></w:pPr>'
        '<w:r><w:fldChar w:fldCharType="begin" w:dirty="true"/>'
        '<w:instrText xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>'
        '<w:fldChar w:fldCharType="separate"/>'
        '<w:t>Обновите содержание клавишей F9 в Word (правый клик → «Обновить поле»).</w:t>'
        '<w:fldChar w:fldCharType="end"/></w:r></w:p>'
    )
    page_break = '<w:p><w:r><w:br w:type="page"/></w:r></w:p>'

    paras = list(_PARA_RE.finditer(document_xml))
    headings = [m for m in paras if _para_heading_level(m.group(0)) == 1]
    if headings and _para_text(headings[0].group(0)).strip().lower() in ("реферат", "аннотация"):
        # СОДЕРЖАНИЕ после реферата: page break до блока; break после не нужен,
        # если следующий раздел уже начинается с новой страницы
        following = [m for m in headings[1:]]
        if following:
            nxt = following[0]
            nxt_xml = nxt.group(0)
            has_own_break = "<w:pageBreakBefore" in nxt_xml
            prev_slice = document_xml[: nxt.start()]
            block = page_break + toc_core + ("" if has_own_break else page_break)
            return prev_slice + block + document_xml[nxt.start():]

    return re.sub(
        r"(<w:body[^>]*>)",
        lambda m: m.group(1) + toc_core + page_break,
        document_xml,
        count=1,
    )


def fill_report_stats(md_path: Path, body_path: Path, final_path: Path) -> bool:
    """Substitute реферат statistics placeholders in the final docx.

    {{FIGURES}}/{{TABLES}} — counted in the built body (all images/tables,
    приложения included, per ГОСТ 7.32). {{SOURCES}} — numbered items of the
    «список использованных источников» section in the Markdown source.
    {{APPENDICES}} — «# Приложение X» headings. {{PAGES}} is left for
    update_fields.py (needs real pagination from LibreOffice).
    Returns True if the document uses any stats placeholder.
    """
    md = md_path.read_text(encoding="utf-8-sig")
    stats_phs = ("{{PAGES}}", "{{FIGURES}}", "{{TABLES}}", "{{SOURCES}}", "{{APPENDICES}}")
    if not any(ph in md for ph in stats_phs):
        return False

    body_xml = zipfile.ZipFile(body_path).read("word/document.xml").decode("utf-8")
    figures = len(re.findall(r"<w:drawing", body_xml))
    tables = len(re.findall(r"<w:tbl>", body_xml))

    sources = 0
    m = re.search(
        r"(?mis)^#\s+список использованных источников\s*$(.*?)(?=^#\s|\Z)", md
    )
    if m:
        sources = len(re.findall(r"(?m)^\s*\d+\.\s", m.group(1)))
    appendices = len(re.findall(r"(?mi)^#\s+приложение\s+[а-яё]\s*$", md))

    values = {
        "{{FIGURES}}": str(figures),
        "{{TABLES}}": str(tables),
        "{{SOURCES}}": str(sources),
        "{{APPENDICES}}": str(appendices),
    }
    with zipfile.ZipFile(final_path, "r") as zf:
        names = zf.namelist()
        files = {n: zf.read(n) for n in names}
    document_xml = files["word/document.xml"].decode("utf-8")
    for ph, value in values.items():
        document_xml = document_xml.replace(ph, value)
    files["word/document.xml"] = document_xml.encode("utf-8")
    final_path.unlink()
    with zipfile.ZipFile(final_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for n, data in files.items():
            zout.writestr(n, data)
    print(f"Report stats: {figures} рис., {tables} табл., {sources} источн., {appendices} прил.")
    return True


def merge_documents(title_path: Path, body_path: Path, output_path: Path):
    title_doc = Document(str(title_path))
    composer = Composer(title_doc)
    body_doc = Document(str(body_path))
    composer.append(body_doc)
    composer.save(str(output_path))


def _find_soffice():
    """LibreOffice: PATH, затем стандартные места Windows/macOS (инсталляторы не прописывают PATH)."""
    found = shutil.which("soffice")
    if found:
        return found
    candidates = [
        Path(r"C:\Program Files\LibreOffice\program\soffice.exe"),
        Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"),
        Path("/Applications/LibreOffice.app/Contents/MacOS/soffice"),
    ]
    for c in candidates:
        if c.exists():
            return str(c)
    return None


def main():
    # Windows-консоль с legacy-кодировкой (cp1252/cp866) падает на кириллице в print
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

    if len(sys.argv) < 2:
        print("Usage: python build.py <report.md> [output.docx]")
        sys.exit(1)

    if not shutil.which("pandoc"):
        print("Error: pandoc not found. Install it: https://pandoc.org/installing.html")
        sys.exit(1)

    md_path = Path(sys.argv[1]).resolve()
    if not md_path.exists():
        print(f"File not found: {md_path}")
        sys.exit(1)

    output_path = (
        Path(sys.argv[2]).resolve() if len(sys.argv) >= 3 else md_path.with_suffix(".docx")
    )

    metadata = parse_yaml_front_matter(md_path)
    print(f"Metadata: {metadata}")

    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir = Path(tmpdir)

        title_template = resolve_title_page(metadata)
        filled = tmpdir / "filled.docx"
        fill_title_page(title_template, metadata, filled)
        transplant_reference_styles(filled)
        print(f"Title page filled ({title_template.stem})")

        body_docx = tmpdir / "body.docx"
        build_report_body(md_path, body_docx)
        _postprocess_body(body_docx, metadata)
        print("Report body built")

        merge_documents(filled, body_docx, output_path)
        stats_used = fill_report_stats(md_path, body_docx, output_path)
        print(f"Report saved: {output_path}")

    add_toc = str(metadata.get("add_toc", "")).strip().lower() in ("true", "yes", "1")
    if add_toc or stats_used:
        soffice = _find_soffice()
        if soffice:
            updater = SKILL_DIR / "scripts" / "update_fields.py"
            try:
                subprocess.run(
                    [sys.executable, str(updater), str(output_path), soffice],
                    check=True,
                    timeout=180,
                )
                print("TOC computed (LibreOffice)")
            except Exception as exc:
                print(f"Warning: TOC not auto-updated ({exc}); откройте в Word и нажмите F9")
        else:
            print(
                "Warning: LibreOffice не найден — СОДЕРЖАНИЕ обновится в Word "
                "(Ctrl+A, F9), а {{PAGES}} в реферате останется плейсхолдером"
            )


if __name__ == "__main__":
    main()
